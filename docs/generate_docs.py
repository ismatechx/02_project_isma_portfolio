from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── colour palette ──────────────────────────────────────────────
TEAL      = RGBColor(0x00, 0xD4, 0xAA)
DARK_BG   = RGBColor(0x06, 0x0D, 0x14)
NAVY      = RGBColor(0x0C, 0x18, 0x22)
TEXT      = RGBColor(0x1A, 0x1A, 0x2E)
MUTED     = RGBColor(0x6B, 0x8F, 0xAF)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF4, 0xF7, 0xFA)
BORDER    = RGBColor(0xD0, 0xDC, 0xE8)

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=20, bold=True, color=TEAL)
        # underline via border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '00D4AA')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        set_font(run, size=14, bold=True, color=TEXT)
    elif level == 3:
        set_font(run, size=12, bold=True, color=MUTED)
    return p

def add_body(doc, text, muted=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=11, color=MUTED if muted else TEXT)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ': ')
        set_font(r1, size=11, bold=True, color=TEXT)
        r2 = p.add_run(text)
        set_font(r2, size=11, color=TEXT)
    else:
        run = p.add_run(text)
        set_font(run, size=11, color=TEXT)

def add_table(doc, headers, rows, shaded_header=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=WHITE)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '00D4AA')
        tcPr.append(shd)

    # data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        fill = 'F4F7FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10, color=TEXT)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

    doc.add_paragraph()
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D0DCE8')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_header_block(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(title + '\n')
    set_font(r1, size=26, bold=True, color=TEAL)
    r2 = p.add_run(subtitle)
    set_font(r2, size=12, color=MUTED, italic=True)
    p.paragraph_format.space_after = Pt(16)

def base_doc():
    doc = Document()
    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    # default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    return doc


# ════════════════════════════════════════════════════════════════
# 1. README.docx
# ════════════════════════════════════════════════════════════════
def make_readme():
    doc = base_doc()
    add_header_block(doc, 'Ismail Uthuman — Portfolio', 'README & Setup Guide')

    add_heading(doc, 'Project Overview')
    add_body(doc, 'Personal portfolio website for Ismail Uthuman, Senior Business Intelligence Analyst with 5+ years of experience in Tableau, Power BI, SQL, and cloud data platforms.')
    add_body(doc, 'Live Site: https://ismatechx.github.io/02_project_isma_portfolio/', muted=True)
    add_body(doc, 'Repository: https://github.com/ismatechx/02_project_isma_portfolio', muted=True)
    add_divider(doc)

    add_heading(doc, 'Features', 2)
    for f in [
        'Animated hero section with typewriter effect and floating data metric cards',
        'Scroll-reveal animations throughout all sections',
        'Responsive mobile navigation with hamburger menu',
        'Experience timeline, skills categories, and certifications sections',
        'Project showcase (3+2 grid) with live Tableau dashboard modal embeds',
        'Supabase-powered contact form — submissions stored in PostgreSQL database',
        'Deployed on GitHub Pages — auto-deploys on every git push',
    ]:
        add_bullet(doc, f)
    add_divider(doc)

    add_heading(doc, 'Tech Stack', 2)
    add_table(doc,
        ['Layer', 'Technology'],
        [
            ('Frontend', 'HTML5, CSS3, Vanilla JavaScript (ES Modules)'),
            ('Typography', 'Google Fonts — Inter'),
            ('Database', 'Supabase (PostgreSQL) — Free Tier'),
            ('Hosting', 'GitHub Pages — Free'),
            ('BI Embeds', 'Tableau Public Embed API'),
            ('Version Control', 'Git / GitHub'),
        ]
    )

    add_heading(doc, 'Local Setup', 2)
    add_body(doc, 'No build tools required. Clone the repo and open index.html directly in a browser, or use a local server:')
    add_bullet(doc, 'VS Code: right-click index.html → Open with Live Server', bold_prefix='Option 1')
    add_bullet(doc, 'python3 -m http.server 8000  →  open http://localhost:8000', bold_prefix='Option 2')
    add_divider(doc)

    add_heading(doc, 'Supabase Contact Form Setup', 2)
    add_heading(doc, 'Step 1 — Create the table', 3)
    add_body(doc, 'Run the following SQL in your Supabase SQL Editor:')
    p = doc.add_paragraph()
    run = p.add_run(
        'create table contact_submissions (\n'
        '  id uuid default gen_random_uuid() primary key,\n'
        '  name text not null,\n'
        '  email text not null,\n'
        '  message text not null,\n'
        '  created_at timestamp with time zone default now()\n'
        ');\n\n'
        'grant usage on schema public to anon;\n'
        'grant insert on public.contact_submissions to anon;'
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT
    p.paragraph_format.left_indent = Inches(0.3)

    add_heading(doc, 'Step 2 — Update credentials in js/main.js', 3)
    add_body(doc, 'Replace YOUR_SUPABASE_PROJECT_URL and YOUR_SUPABASE_ANON_KEY with your project values found in Supabase Dashboard → Settings → API.')
    add_divider(doc)

    add_heading(doc, 'Deployment', 2)
    add_body(doc, 'The site auto-deploys to GitHub Pages on every push to the main branch.')
    add_table(doc,
        ['Step', 'Command'],
        [
            ('Stage changes', 'git add .'),
            ('Commit', 'git commit -m "your message"'),
            ('Push & deploy', 'git push origin main'),
        ]
    )

    doc.save(os.path.join(OUTPUT_DIR, 'README.docx'))
    print('✓ README.docx')


# ════════════════════════════════════════════════════════════════
# 2. DESIGN.docx
# ════════════════════════════════════════════════════════════════
def make_design():
    doc = base_doc()
    add_header_block(doc, 'Design Document', 'Ismail Uthuman Portfolio — UI/UX Specification')

    add_heading(doc, 'Design Philosophy')
    add_table(doc,
        ['Principle', 'Application'],
        [
            ('Data-inspired', 'Floating KPI metric cards in hero, stat counters, chart-like gradient banners'),
            ('Dark & premium', 'Deep navy/charcoal backgrounds with teal accent — professional and modern'),
            ('Minimal noise', 'Muted secondary text, restrained colour use, clean whitespace'),
            ('Progressive disclosure', 'Scroll-reveal animations guide users through content naturally'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Colour Palette')
    add_table(doc,
        ['Token', 'Hex Value', 'Usage'],
        [
            ('--clr-bg', '#060d14', 'Page background'),
            ('--clr-surface', '#0c1822', 'Section backgrounds, cards'),
            ('--clr-surface-2', '#111f2e', 'Elevated cards, form inputs'),
            ('--clr-accent', '#00d4aa', 'Primary accent — teal/green'),
            ('--clr-accent-hover', '#00b891', 'Hover state for accent elements'),
            ('--clr-accent-glow', 'rgba(0,212,170,0.15)', 'Glow effects, hover fills'),
            ('--clr-text', '#e2eeff', 'Primary body text'),
            ('--clr-muted', '#6b8faf', 'Secondary/supporting text'),
            ('--clr-border', '#162536', 'Borders and dividers'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Typography')
    add_table(doc,
        ['Element', 'Font', 'Size', 'Weight'],
        [
            ('Font Family', 'Inter (Google Fonts)', '—', '—'),
            ('Hero H1', 'Inter', 'clamp(2.8rem, 7vw, 5.5rem)', '800'),
            ('Section Titles', 'Inter', '1.9rem', '800'),
            ('Card Headings', 'Inter', '0.92rem', '700'),
            ('Body Text', 'Inter', '0.88–0.97rem', '400'),
            ('Labels / Tags', 'Inter', '0.68–0.75rem', '600–700'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Layout & Spacing')
    add_bullet(doc, '1080px, centred, 2rem horizontal padding', bold_prefix='Container max-width')
    add_bullet(doc, '7rem top and bottom per section', bold_prefix='Section padding')
    add_bullet(doc, '--radius: 10px   --radius-lg: 16px', bold_prefix='Border radius')
    add_bullet(doc, '1.25–1.5rem between grid items', bold_prefix='Grid gaps')
    add_divider(doc)

    add_heading(doc, 'Section Layouts')
    add_table(doc,
        ['Section', 'Layout'],
        [
            ('Hero', 'Full viewport height, centred content, 3 animated orbs, 4 floating data cards'),
            ('About', '2-column grid: 220×220px profile photo + bio text + 3 stat counters'),
            ('Skills', 'Stacked groups by category — pill-shaped tags with hover glow'),
            ('Experience', 'Vertical timeline with gradient line, dot markers, date badges'),
            ('Projects', 'Row 1: 3 equal cards (span 2 each) | Row 2: 2 wider cards (span 3 each)'),
            ('Certifications', '2-column grid: Certifications (left) | Awards (right)'),
            ('Contact', '2-column: contact links (left) | contact form (right)'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Animations')
    add_table(doc,
        ['Animation', 'Element', 'Duration', 'Behaviour'],
        [
            ('orbFloat', 'Hero background orbs', '10–16s', 'Float + scale, alternate, infinite'),
            ('floatCard', 'Hero metric cards', '5s', 'Vertical float, staggered delays'),
            ('pulse', 'Hero badge dot', '2s', 'Opacity + scale + box-shadow glow'),
            ('blink', 'Typewriter cursor', '0.75s', 'Step-end opacity blink'),
            ('bounce', 'Scroll indicator', '2.5s', 'translateY bounce'),
            ('taxiShimmer', 'NYC Taxi card banner', '2.5s', 'Diagonal gold shimmer sweep'),
            ('taxiGlow', 'NYC Taxi card border', '2.5s', 'Border glow pulse, alternate'),
            ('Scroll reveal', '.reveal elements', '0.6s', 'translateY(24px)→0 + fade, staggered 80ms'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Responsive Breakpoints')
    add_table(doc,
        ['Breakpoint', 'Changes Applied'],
        [
            ('max-width: 900px', 'Floating hero data cards hidden'),
            ('max-width: 768px', 'Mobile nav (hamburger menu), single-column layouts for About / Certs / Contact / Projects'),
        ]
    )

    doc.save(os.path.join(OUTPUT_DIR, 'DESIGN.docx'))
    print('✓ DESIGN.docx')


# ════════════════════════════════════════════════════════════════
# 3. TECHNICAL_SPEC.docx
# ════════════════════════════════════════════════════════════════
def make_technical():
    doc = base_doc()
    add_header_block(doc, 'Technical Specification', 'Ismail Uthuman Portfolio — Architecture & Integration Guide')

    add_heading(doc, 'Project Overview')
    add_table(doc,
        ['Property', 'Detail'],
        [
            ('Type', 'Static single-page portfolio website'),
            ('Owner', 'Ismail Uthuman'),
            ('Live URL', 'https://ismatechx.github.io/02_project_isma_portfolio/'),
            ('Repository', 'https://github.com/ismatechx/02_project_isma_portfolio'),
            ('Last Updated', 'February 2026'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Architecture')
    add_body(doc, 'The site is a zero-dependency static web app. No build tools, no bundler, no framework. JavaScript uses ES Modules for the Supabase import.')
    add_body(doc, 'Browser  →  GitHub Pages (static hosting)  →  index.html + CSS + JS', muted=True)
    add_body(doc, '                                              └── Supabase REST API (contact form)', muted=True)
    add_body(doc, '                                              └── Tableau Public Embed API (dashboards)', muted=True)
    add_divider(doc)

    add_heading(doc, 'Tech Stack')
    add_heading(doc, 'Frontend', 3)
    add_table(doc,
        ['Technology', 'Version', 'Purpose'],
        [
            ('HTML5', '—', 'Page structure and semantics'),
            ('CSS3', '—', 'Styling, animations, responsive layout'),
            ('JavaScript', 'ES2020+', 'Interactivity, Supabase client, Tableau loader'),
            ('Google Fonts — Inter', '—', 'Typography'),
        ]
    )
    add_heading(doc, 'Backend / Services', 3)
    add_table(doc,
        ['Service', 'Plan', 'Purpose'],
        [
            ('Supabase', 'Free', 'PostgreSQL database for contact form submissions'),
            ('GitHub Pages', 'Free', 'Static hosting, auto-deploy from main branch'),
            ('Tableau Public', 'Free', 'Embedded interactive dashboards'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'JavaScript Modules (main.js)')
    add_table(doc,
        ['Module', 'Functionality'],
        [
            ('Nav', 'Scroll adds scrolled class to header; active link tracking via section offsets'),
            ('Mobile Nav', 'Hamburger toggle; close menu on link click'),
            ('Typewriter', 'Cycles through role titles with typing/deleting animation'),
            ('Scroll Reveal', 'IntersectionObserver on .reveal elements — staggered fade-up on enter'),
            ('Tableau Modals', 'Lazy-loads Tableau embed script on first open; supports multiple dashboards via data-viz'),
            ('Contact Form', 'Async submit to Supabase; UI state: Sending → Sent / Error'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Supabase Integration')
    add_heading(doc, 'Database Schema — contact_submissions', 3)
    add_table(doc,
        ['Column', 'Type', 'Constraints'],
        [
            ('id', 'uuid', 'Primary key, gen_random_uuid()'),
            ('name', 'text', 'NOT NULL'),
            ('email', 'text', 'NOT NULL'),
            ('message', 'text', 'NOT NULL'),
            ('created_at', 'timestamptz', 'Default: now()'),
        ]
    )
    add_heading(doc, 'Security', 3)
    add_table(doc,
        ['Setting', 'Value'],
        [
            ('Row Level Security', 'Disabled on contact_submissions (insert-only public table)'),
            ('Permissions granted', 'anon role: INSERT on contact_submissions'),
            ('Key used in frontend', 'Publishable (anon) key only — safe to expose in browser'),
            ('Service role key', 'Never used in frontend code'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Tableau Embed Integration')
    add_heading(doc, 'Supported Dashboards', 3)
    add_table(doc,
        ['#', 'Dashboard', 'Viz ID'],
        [
            ('1', 'NYC Maven Taxi Challenge', 'viz1771970012538'),
            ('2', 'World Happiness Report 2022', 'viz1771976526694'),
        ]
    )
    add_heading(doc, 'Behaviour', 3)
    add_bullet(doc, 'Triggered by .tableau-modal-btn with data-viz attribute')
    add_bullet(doc, 'Tableau JS API script injected lazily — only on first modal open per dashboard')
    add_bullet(doc, 'Closed by: close button, overlay click, or Escape key')
    add_divider(doc)

    add_heading(doc, 'Performance & Browser Support')
    add_table(doc,
        ['Technique / Browser', 'Detail'],
        [
            ('Lazy Tableau loading', 'Tableau JS API only loads when modal is opened'),
            ('CSS cache busting', '?v=N query param on stylesheet link'),
            ('Zero bundle', 'Pure HTML/CSS/JS — no framework overhead'),
            ('Chrome 90+', 'Full support'),
            ('Firefox 88+', 'Full support'),
            ('Safari 14+', 'Full support (backdrop-filter supported)'),
            ('Edge 90+', 'Full support'),
            ('Mobile', 'Responsive — floating hero cards hidden on small screens'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Known Limitations')
    add_table(doc,
        ['Issue', 'Detail'],
        [
            ('Contact form RLS', 'sb_publishable_ key format does not map to anon RLS role — RLS disabled on table'),
            ('Email notifications', 'Not yet implemented — form stores to DB only, no email sent to owner'),
            ('Spam protection', 'No CAPTCHA or rate limiting on contact form currently'),
        ]
    )

    doc.save(os.path.join(OUTPUT_DIR, 'TECHNICAL_SPEC.docx'))
    print('✓ TECHNICAL_SPEC.docx')


# ════════════════════════════════════════════════════════════════
# 4. PROJECT_BRIEF.docx
# ════════════════════════════════════════════════════════════════
def make_brief():
    doc = base_doc()
    add_header_block(doc, 'Project Brief', 'Ismail Uthuman Portfolio Website')

    add_heading(doc, 'Summary')
    add_body(doc,
        'A fully custom, live portfolio website built to showcase Ismail Uthuman\'s expertise as a '
        'Senior Business Intelligence Analyst. Hosted on GitHub Pages, connected to a real Supabase '
        'database, and featuring interactive embedded Tableau dashboards — demonstrating both BI skills '
        'and full-stack technical capability.'
    )
    add_divider(doc)

    add_heading(doc, 'Problem Statement')
    add_body(doc, 'Off-the-shelf portfolio platforms do not provide:')
    for b in [
        'Full control over design, branding, and layout',
        'Ability to embed live, interactive Tableau dashboards',
        'A functional contact form backed by a real database',
        'A single shareable link for job applications and networking',
    ]:
        add_bullet(doc, b)
    add_divider(doc)

    add_heading(doc, 'Solution')
    add_table(doc,
        ['What', 'How'],
        [
            ('Custom dark UI', 'Pure HTML/CSS with teal accent — no templates used'),
            ('Live hosting', 'GitHub Pages — auto-deploys on every update to main branch'),
            ('Contact form with DB', 'Supabase (PostgreSQL) — all messages stored and queryable'),
            ('Tableau dashboards', 'Embedded live via Tableau Public Embed API in modal overlays'),
            ('Animations & UX', 'CSS keyframes + IntersectionObserver scroll reveal'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Projects Showcased')
    add_table(doc,
        ['#', 'Project', 'Type', 'Tech Stack'],
        [
            ('01', 'Supply Operations Dashboard', 'Enterprise — Zuci Systems', 'Tableau, SQL, Tableau Prep'),
            ('02', 'Global Fund Health Dashboards', 'Enterprise — Radiare', 'Tableau, GraphQL, Tableau Server'),
            ('03', 'ISMS Risk Dashboard', 'Enterprise — Zuci Systems', 'Power BI, RLS, MS SQL Server'),
            ('04', 'NYC Maven Taxi Challenge', 'Public — Live Embed', 'Tableau Public'),
            ('05', 'World Happiness Report 2022', 'Public — Live Embed', 'Tableau Public'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Results & Metrics')
    add_table(doc,
        ['Metric', 'Value'],
        [
            ('Hosting cost', '$0 / month'),
            ('Database cost', '$0 / month (Supabase free tier)'),
            ('Deployment time', '~30 seconds after git push'),
            ('Mobile responsive', 'Yes — tested across breakpoints'),
            ('Live Tableau dashboards', '2 embedded (NYC Taxi + World Happiness)'),
            ('Contact form', 'Fully functional with Supabase DB storage'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Links')
    add_table(doc,
        ['Resource', 'URL'],
        [
            ('Live Site', 'https://ismatechx.github.io/02_project_isma_portfolio/'),
            ('GitHub Repo', 'https://github.com/ismatechx/02_project_isma_portfolio'),
            ('Tableau Public', 'https://public.tableau.com/app/profile/ismail4056'),
            ('LinkedIn', 'https://www.linkedin.com/in/isma96u/'),
        ]
    )
    add_divider(doc)

    add_heading(doc, 'Author')
    add_body(doc, 'Ismail Uthuman')
    add_body(doc, 'Senior Business Intelligence Analyst', muted=True)
    add_body(doc, 'isma96.u@gmail.com  ·  New Jersey, USA', muted=True)

    doc.save(os.path.join(OUTPUT_DIR, 'PROJECT_BRIEF.docx'))
    print('✓ PROJECT_BRIEF.docx')


if __name__ == '__main__':
    make_readme()
    make_design()
    make_technical()
    make_brief()
    print('\nAll 4 Word documents created in docs/')
